from docx import Document

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.shared import Cm

import sys
# 导入必要的库
from docx import Document  # 用于操作Word文档
import matplotlib  # 绘图基础库
import ctypes  # 用于调用C/C++共享内存
from ctypes import wintypes  # Windows类型定义
import time  # 时间相关操作

# 设置中文字体（解决乱码）
matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体显示中文
matplotlib.rcParams['axes.unicode_minus'] = False    # 解决负号显示问题

# 定义Windows API常量
PAGE_READWRITE = 0x04  # 内存读写权限
FILE_MAP_READ = 0x0004  # 文件映射读取权限
INVALID_HANDLE_VALUE = -1  # 无效句柄值

# 加载Windows内核库
kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)

# 定义Windows API函数原型
OpenFileMapping = kernel32.OpenFileMappingW
OpenFileMapping.argtypes = [wintypes.DWORD, wintypes.BOOL, wintypes.LPCWSTR]
OpenFileMapping.restype = wintypes.HANDLE

MapViewOfFile = kernel32.MapViewOfFile
MapViewOfFile.argtypes = [wintypes.HANDLE, wintypes.DWORD,
                          wintypes.DWORD, wintypes.DWORD, ctypes.c_size_t]
MapViewOfFile.restype = wintypes.LPVOID

UnmapViewOfFile = kernel32.UnmapViewOfFile
UnmapViewOfFile.argtypes = [wintypes.LPCVOID]
UnmapViewOfFile.restype = wintypes.BOOL

CloseHandle = kernel32.CloseHandle
CloseHandle.argtypes = [wintypes.HANDLE]
CloseHandle.restype = wintypes.BOOL




def read_single_string_infinite(map_name="Local\\IllustrateMemory", retry_interval=1, max_size=256):
    """
    无限重试方式，从共享内存读取一个字符串。
    假设内存布局：
    前4字节：int，字符串字节长度
    后续字节：UTF-8编码的字符串数据
    
    参数:
        map_name: 共享内存名称
        retry_interval: 重试间隔秒
        max_size: 共享内存最大字节数（包含长度4字节）

    返回:
        解码后的字符串
    """
    import ctypes, time, struct

    size = max_size
    hMap = None

    print(f"尝试连接共享内存 '{map_name}'（无限重试中）...")
    while True:
        hMap = OpenFileMapping(FILE_MAP_READ, False, map_name)
        if hMap:
            break
        time.sleep(retry_interval)

    pBuf = MapViewOfFile(hMap, FILE_MAP_READ, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        raw_bytes = bytes(buffer[:])

        # 读取前4字节字符串长度
        str_len = struct.unpack('i', raw_bytes[:4])[0]
        if str_len < 0 or str_len > size - 4:
            raise ValueError("共享内存中字符串长度异常")

        # 读取字符串字节
        str_bytes = raw_bytes[4:4+str_len]

        # 解码成字符串
        return str_bytes.decode('utf-8')

    finally:
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)

FILE_MAP_READ = 0x0004
FILE_MAP_WRITE = 0x0002
FILE_MAP_ALL_ACCESS = 0x001f
def clear_shared_memory(map_name="Local\\IllustrateMemory", max_size=256):
    """
    清空共享内存，将内容全部写为0。
    """
    size = max_size

    # 打开共享内存，需写权限
    hMap = OpenFileMapping(FILE_MAP_WRITE, False, map_name)
    if not hMap:
        raise ctypes.WinError(ctypes.get_last_error())

    pBuf = MapViewOfFile(hMap, FILE_MAP_WRITE, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        # 将共享内存全部置0
        for i in range(size):
            buffer[i] = b'\x00'
    finally:
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)

# val = read_single_string_infinite("Local\\IllustrateMemory")
while True:
    val = read_single_string_infinite("Local\\IllustrateMemory")
    if val.strip() != "":
        break
    time.sleep(0.5)
clear_shared_memory()
print("读取到的 int 值是：", val)

from docx.enum.text import WD_ALIGN_PARAGRAPH  # 添加这一行
def insert_picture_after_bookmark(doc, bookmark_name, image_path):
    # 找到bookmarkStart元素
    bookmark_elem = None
    for child in doc.element.iter():
        if child.tag == qn('w:bookmarkStart') and child.get(qn('w:name')) == bookmark_name:
            bookmark_elem = child
            break
    if bookmark_elem is None:
        raise ValueError(f"书签 '{bookmark_name}' 未找到")

    # 获取书签所在段落节点（bookmarkElem的祖先w:p）
    para_elem = bookmark_elem
    while para_elem.tag != qn('w:p'):
        para_elem = para_elem.getparent()
        if para_elem is None:
            raise RuntimeError("没找到对应段落")

    # 在段落后插入一个新段落
    new_p = OxmlElement('w:p')
    para_elem.addnext(new_p)

    # 用 python-docx 的 Document 对象包装该段落方便插入图片
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, doc)
    new_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # 添加这一行：设置右对齐

    new_para.add_run().add_picture(image_path, width=Cm(6.77), height=Cm(6.77))

    return doc




if val:
    doc = Document(val)
    insert_picture_after_bookmark(doc, "ChartPlaceholder", "output_image.png")
    doc.save(val)

sys.exit(0)  # 退出Python程序，0表示正常退出

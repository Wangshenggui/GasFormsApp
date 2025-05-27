# 导入必要的库
from docx import Document  # 用于操作Word文档
from docx.shared import Inches  # 用于设置图片尺寸
import matplotlib.pyplot as plt  # 绘图库
import numpy as np  # 数值计算库
import io  # 用于内存字节流操作
import matplotlib  # 绘图基础库
import os  # 操作系统接口
from PIL import Image  # 图像处理库
import win32clipboard  # Windows剪贴板操作
from win32con import CF_DIB  # 剪贴板格式常量
import ctypes  # 用于调用C/C++共享内存
from ctypes import wintypes  # Windows类型定义
import struct  # 字节打包/解包
import time  # 时间相关操作

# 设置中文字体（解决乱码）
matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体显示中文
matplotlib.rcParams['axes.unicode_minus'] = False    # 解决负号显示问题

# 定义Windows API常量
PAGE_READWRITE = 0x04  # 内存读写权限
FILE_MAP_READ = 0x0004  # 文件映射读取权限
INVALID_HANDLE_VALUE = -1  # 无效句柄值

# 加载Windows内核库
kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)

# 定义Windows API函数原型
OpenFileMapping = kernel32.OpenFileMappingW
OpenFileMapping.argtypes = [wintypes.DWORD, wintypes.BOOL, wintypes.LPCWSTR]
OpenFileMapping.restype = wintypes.HANDLE

MapViewOfFile = kernel32.MapViewOfFile
MapViewOfFile.argtypes = [wintypes.HANDLE, wintypes.DWORD,
                          wintypes.DWORD, wintypes.DWORD, ctypes.c_size_t]
MapViewOfFile.restype = wintypes.LPVOID

UnmapViewOfFile = kernel32.UnmapViewOfFile
UnmapViewOfFile.argtypes = [wintypes.LPCVOID]
UnmapViewOfFile.restype = wintypes.BOOL

CloseHandle = kernel32.CloseHandle
CloseHandle.argtypes = [wintypes.HANDLE]
CloseHandle.restype = wintypes.BOOL

def read_shared_memory(map_name="Local\\MySharedMemory", rows=60, cols=2, retries=10, retry_interval=1):
    """
    从共享内存读取数据
    
    参数:
        map_name: 共享内存名称
        rows: 数据行数
        cols: 数据列数
        retries: 重试次数
        retry_interval: 重试间隔(秒)
    
    返回:
        二维列表形式的数据矩阵
    """
    size = rows * cols * 8  # 计算内存大小(每个double占8字节)
    hMap = None  # 内存映射句柄

    # 重试机制打开共享内存
    for attempt in range(retries):
        hMap = OpenFileMapping(FILE_MAP_READ, False, map_name)
        if hMap:
            break
        time.sleep(retry_interval)

    if not hMap:
        raise RuntimeError("无法打开共享内存")

    # 映射内存视图
    pBuf = MapViewOfFile(hMap, FILE_MAP_READ, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        # 读取内存数据并转换为Python数据结构
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        read_bytes = bytes(buffer[:])
        values = struct.unpack('d' * rows * cols, read_bytes)
        matrix = [values[i*cols:(i+1)*cols] for i in range(rows)]
        return matrix
    finally:
        # 清理资源
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)

# 从共享内存读取数据
data = read_shared_memory()[:12]
# 先筛选出满足条件的点
x = np.array([row[0] for row in data if row[0] != 0 and row[1] != 0])
y = np.array([row[1] for row in data if row[0] != 0 and row[1] != 0])
# 线性拟合(1次多项式拟合)
coefficients = np.polyfit(x, y, 1)
slope, intercept = coefficients  # 获取斜率和截距
fitted_y = slope * x + intercept  # 计算拟合值
# 计算R平方值(拟合优度)
ss_res = np.sum((y - fitted_y) ** 2)  # 残差平方和
ss_tot = np.sum((y - np.mean(y)) ** 2)  # 总平方和
r_squared = 1 - ss_res / ss_tot  # R平方值
# 27.1070993058854
# -40.3230351033197
# 0.998135143421265

# 27.7873781239904
# -41.8748868048842
# 0.997957763956664

# 27.78919622571
# -41.8791481700397
# 0.998634884844042

# 26.8380515458133
# -39.5923622882033
# 0.995913670949004

# 9
# 27.0498955091897
# -40.1140555547572
# 0.996838309000741

# 10
# 27.2160627030406
# -40.5326661485488
# 0.99747741604489

# 11
# 26.901479298497
# -39.7228752854571
# 0.997450870863665

# 12
# 26.8041958696157
# -39.4672491714229
# 0.997871011206052

data = read_shared_memory()[:20]  # 多取一点，防止筛选后点太少

# 筛选满足条件的点（非零）
filtered_points = [row for row in data if row[0] != 0 and row[1] != 0]

best_r2 = -np.inf
best_range = None
best_coefficients = None
DataExtentEnd = 0

# 确保筛选后的点数够12个
max_len = min(len(filtered_points), 12)

for end in range(5, max_len + 1):  # 从第5个点到第max_len个点
    subset = filtered_points[:end]
    x = np.array([row[0] for row in subset])
    y = np.array([row[1] for row in subset])
    coefficients = np.polyfit(x, y, 1)
    slope, intercept = coefficients
    fitted_y = slope * x + intercept
    ss_res = np.sum((y - fitted_y) ** 2)
    ss_tot = np.sum((y - np.mean(y)) ** 2)
    r_squared = 1 - ss_res / ss_tot

    if r_squared > best_r2:
        best_r2 = r_squared
        best_range = (1, end)
        DataExtentEnd = end
        best_coefficients = coefficients



# # 转换共享内存数据为x,y坐标数组(过滤零值)
# x = np.array([row[0] for row in data if row[0] != 0 and row[1] != 0])
# y = np.array([row[1] for row in data if row[0] != 0 and row[1] != 0])

# # 线性拟合(1次多项式拟合)
# coefficients = np.polyfit(x, y, 1)
# slope, intercept = coefficients  # 获取斜率和截距
# fitted_y = slope * x + intercept  # 计算拟合值

# # 计算R平方值(拟合优度)
# ss_res = np.sum((y - fitted_y) ** 2)  # 残差平方和
# ss_tot = np.sum((y - np.mean(y)) ** 2)  # 总平方和
# r_squared = 1 - ss_res / ss_tot  # R平方值

# 从共享内存读取数据
data = read_shared_memory()
# 先筛选出满足条件的点
x = np.array([row[0] for row in data if row[0] != 0 and row[1] != 0])
y = np.array([row[1] for row in data if row[0] != 0 and row[1] != 0])
# 创建图形和坐标轴
fig, ax = plt.subplots(figsize=(6, 6))

# 绘制散点图(原始数据)和拟合线
ax.scatter(x, y, s=4, color='black', label='原始数据')
data = read_shared_memory()[:DataExtentEnd]
# 先筛选出满足条件的点
x = np.array([row[0] for row in data if row[0] != 0 and row[1] != 0])
y = np.array([row[1] for row in data if row[0] != 0 and row[1] != 0])
ax.scatter(x, y, s=10, color='black', marker='s',label='原始数据')
slope = best_coefficients[0]
intercept = best_coefficients[1]
x_new = np.linspace(0, np.max(x)+1, 100)
y_new = slope * x_new + intercept
ax.plot(x_new, y_new, color='black', label=f'拟合线: y = {slope:.3f}x + {intercept:.3f}')

sign = '+' if intercept >= 0 else '-'
abs_intercept = abs(intercept)
ax.text(0.6, 275, f'y = {slope:.3f}x {sign} {abs_intercept:.3f} \n    R2 = {best_r2:.5f}', fontsize=18, color='black')


# # 设置图表标题
# sign = '+' if intercept >= 0 else '-'
# abs_intercept = abs(intercept)
# ax.set_title(f'y = {slope:.3f}x {sign} {abs_intercept:.3f} \nR2 = {best_r2:.5f}', pad=15,fontsize=20)

# 设置坐标轴范围
ax.set_xlim(0, 9)
ax.set_ylim(-300, 320)

# 设置刻度
ax.set_xticks(np.arange(1, 9, 1))  # X轴1到8，步长1
ax.set_yticks(np.arange(-300, 301, 50))  # Y轴-60到140，步长20
ax.tick_params(axis='x', labelsize=20)  # 设置X轴刻度标签字体大小为12
ax.tick_params(axis='y', labelsize=20)  # 设置Y轴刻度标签字体大小为12


# 设置坐标轴样式
# ax.spines['top'].set_visible(False)  # 隐藏上边框
# ax.spines['right'].set_visible(False)  # 隐藏右边框
ax.spines['left'].set_position('zero')  # 左边框移动到零位置
ax.spines['bottom'].set_position('zero')  # 下边框移动到零位置
# 隐藏默认坐标轴线
for spine in ax.spines.values():
    spine.set_visible(False)

# 画一条自定义的X轴线，长度比坐标轴范围短一点
ax.plot([0, 7.9], [0, 0], color='black', lw=1.0)    # X轴线：水平线，从x=0到7.9，y=0固定
ax.plot([0, 0], [-300, 300], color='black', lw=1.0)  # Y轴线：竖直线，x=0固定，y从-60到150


# 添加箭头风格的坐标轴
ax.annotate('', xy=(8.5, 0), xytext=(0, 0),
            arrowprops=dict(arrowstyle='->', lw=1.0, color='black'))  # X轴箭头
ax.annotate('', xy=(0, 320), xytext=(0, -302),
            arrowprops=dict(arrowstyle='->', lw=1.0, color='black'))  # Y轴箭头

# 保存图像到内存字节流
img_bytes = io.BytesIO()
plt.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
img_bytes.seek(0)  # 将指针移回起始位置

# 把图像插入 Word 文档
doc = Document()
doc.add_heading('去除X轴0刻度的图', level=1)
doc.add_paragraph('图中X轴从0开始，但不显示0刻度，仍保留箭头和线性拟合线：')
doc.add_picture(img_bytes, width=Inches(5.0))
doc.save('你干嘛.docx')
# 打开 Word 文档
os.startfile('你干嘛.docx')

# 关闭图形释放内存
plt.close()

# ---------- 复制图像到剪贴板（仅限 Windows） ----------
def send_to_clipboard(im):
    """
    将图像复制到Windows剪贴板
    
    参数:
        im: PIL图像对象
    """
    output = io.BytesIO()
    im.convert('RGB').save(output, 'BMP')  # 转换为BMP格式
    data = output.getvalue()[14:]  # 去掉BMP头的前14字节
    output.close()

    # 操作剪贴板
    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardData(CF_DIB, data)  # 设置剪贴板数据
    win32clipboard.CloseClipboard()

# 从字节流加载为PIL图像并复制到剪贴板
image = Image.open(img_bytes)
send_to_clipboard(image)

# 将拟合结果写入共享内存
import mmap
slope_mem_name = "Local\\tempSharedMemory"  # 共享内存名称
# 打包斜率和截距、R平方值(3个双精度浮点数)
slope_bytes = struct.pack('3d', slope, intercept, r_squared)  

# 创建内存映射并写入数据
with mmap.mmap(-1, len(slope_bytes), tagname=slope_mem_name, access=mmap.ACCESS_WRITE) as slope_mmap:
    slope_mmap.write(slope_bytes)
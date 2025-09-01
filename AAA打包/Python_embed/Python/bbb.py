from docx import Document

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docx.shared import Cm

import sys
# 导入必要的库
from docx import Document  # 用于操作Word文档
import matplotlib  # 绘图基础库
import ctypes  # 用于调用C/C++共享内存
from ctypes import wintypes  # Windows类型定义
import time  # 时间相关操作

# 设置中文字体（解决乱码）
matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 使用黑体显示中文
matplotlib.rcParams['axes.unicode_minus'] = False    # 解决负号显示问题

# 定义Windows API常量
PAGE_READWRITE = 0x04  # 内存读写权限
FILE_MAP_READ = 0x0004  # 文件映射读取权限
INVALID_HANDLE_VALUE = -1  # 无效句柄值

# 加载Windows内核库
kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)

# 定义Windows API函数原型
OpenFileMapping = kernel32.OpenFileMappingW
OpenFileMapping.argtypes = [wintypes.DWORD, wintypes.BOOL, wintypes.LPCWSTR]
OpenFileMapping.restype = wintypes.HANDLE

MapViewOfFile = kernel32.MapViewOfFile
MapViewOfFile.argtypes = [wintypes.HANDLE, wintypes.DWORD,
                          wintypes.DWORD, wintypes.DWORD, ctypes.c_size_t]
MapViewOfFile.restype = wintypes.LPVOID

UnmapViewOfFile = kernel32.UnmapViewOfFile
UnmapViewOfFile.argtypes = [wintypes.LPCVOID]
UnmapViewOfFile.restype = wintypes.BOOL

CloseHandle = kernel32.CloseHandle
CloseHandle.argtypes = [wintypes.HANDLE]
CloseHandle.restype = wintypes.BOOL




def read_single_string_infinite(map_name="Local\\IllustrateMemory", retry_interval=1, max_size=256):
    """
    无限重试方式，从共享内存读取一个字符串。
    假设内存布局：
    前4字节：int，字符串字节长度
    后续字节：UTF-8编码的字符串数据
    
    参数:
        map_name: 共享内存名称
        retry_interval: 重试间隔秒
        max_size: 共享内存最大字节数（包含长度4字节）

    返回:
        解码后的字符串
    """
    import ctypes, time, struct

    size = max_size
    hMap = None

    print(f"尝试连接共享内存 '{map_name}'（无限重试中）...")
    while True:
        hMap = OpenFileMapping(FILE_MAP_READ, False, map_name)
        if hMap:
            break
        time.sleep(retry_interval)

    pBuf = MapViewOfFile(hMap, FILE_MAP_READ, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        raw_bytes = bytes(buffer[:])

        # 读取前4字节字符串长度
        str_len = struct.unpack('i', raw_bytes[:4])[0]
        if str_len < 0 or str_len > size - 4:
            raise ValueError("共享内存中字符串长度异常")

        # 读取字符串字节
        str_bytes = raw_bytes[4:4+str_len]

        # 解码成字符串
        return str_bytes.decode('utf-8')

    finally:
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)

FILE_MAP_READ = 0x0004
FILE_MAP_WRITE = 0x0002
FILE_MAP_ALL_ACCESS = 0x001f
def clear_shared_memory(map_name="Local\\IllustrateMemory", max_size=256):
    """
    清空共享内存，将内容全部写为0。
    """
    size = max_size

    # 打开共享内存，需写权限
    hMap = OpenFileMapping(FILE_MAP_WRITE, False, map_name)
    if not hMap:
        raise ctypes.WinError(ctypes.get_last_error())

    pBuf = MapViewOfFile(hMap, FILE_MAP_WRITE, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        # 将共享内存全部置0
        for i in range(size):
            buffer[i] = b'\x00'
    finally:
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)

# val = read_single_string_infinite("Local\\IllustrateMemory")
while True:
    val = read_single_string_infinite("Local\\IllustrateMemory")
    if val.strip() != "":
        break
    time.sleep(0.5)
clear_shared_memory()
print("读取到的 int 值是：", val)

from docx.enum.text import WD_ALIGN_PARAGRAPH  # 添加这一行
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from PIL import Image
import uuid
import os
import tempfile

def insert_picture_after_bookmark(doc, bookmark_name, image_path):
    print("[INFO] 查找书签...")
    bookmark_elem = None
    for child in doc.element.iter():
        if child.tag == qn('w:bookmarkStart') and child.get(qn('w:name')) == bookmark_name:
            bookmark_elem = child
            break
    if bookmark_elem is None:
        raise ValueError(f"书签 '{bookmark_name}' 未找到")

    para_elem = bookmark_elem
    while para_elem is not None and para_elem.tag != qn('w:p'):
        para_elem = para_elem.getparent()
    if para_elem is None:
        raise RuntimeError("找不到书签所在段落")

    new_p = OxmlElement('w:p')
    para_elem.addnext(new_p)
    new_para = Paragraph(new_p, doc)
    new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    print("[INFO] 正在裁剪图片...")
    a = 0.05  # 裁剪右边 5%
    b = 0.02  # 裁剪左边 2%

    with Image.open(image_path) as img:
        width, height = img.size
        left = int(width * b)
        right = int(width * (1 - a))
        cropped = img.crop((left, 0, right, height))

        # 用 NamedTemporaryFile(delete=False)
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_file:
            temp_path = tmp_file.name
            cropped.save(temp_path)
            print(f"[INFO] 临时文件已保存: {temp_path}")

    try:
        new_para.add_run().add_picture(temp_path, width=Cm(6.77 * (1 - a)), height=Cm(6.77))
        print("[INFO] 图片插入成功")
    finally:
        # 加个重试删除
        for _ in range(3):
            try:
                os.remove(temp_path)
                print("[INFO] 已删除临时文件")
                break
            except PermissionError:
                import time
                time.sleep(0.5)

    return doc


import os

print("当前工作目录:", os.getcwd())
print("脚本所在目录:", os.path.dirname(os.path.abspath(sys.argv[0])))
print("图片路径是否存在:", os.path.exists("output_image.png"))

if val:
    # 切换工作目录到脚本所在目录
    script_dir = os.path.dirname(os.path.abspath(sys.argv[0]))
    os.chdir(script_dir)
    
    try:
        # 获取当前用户的 AppData\Roaming 路径
        appdata_dir = os.path.join(os.environ["APPDATA"], "瓦斯含量测定数据分析系统")
        print(f"[INFO] AppData 路径: {appdata_dir}")

        # 拼接图片路径
        image_path = os.path.join(appdata_dir, "Image", "output_image.png")
        print(f"[INFO] 图片路径: {image_path}")

        # 加载 Word 文档
        print(f"[INFO] 正在加载 Word 文档: {val}")
        doc = Document(val)

        # 调用插入图片的函数
        print(f"[INFO] 调用 insert_picture_after_bookmark 插入图片到书签 'ChartPlaceholder'")
        insert_picture_after_bookmark(doc, "ChartPlaceholder", image_path)

        # 保存文档
        doc.save(val)
        print(f"[INFO] 成功保存文档: {val}")

    except Exception as e:
        print(f"[ERROR] 发生异常: {e}")

sys.exit(0)  # 退出Python程序，0表示正常退出

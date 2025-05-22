from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import io
import matplotlib
import os
from PIL import Image
import win32clipboard
from win32con import CF_DIB
import ctypes
from ctypes import wintypes
import struct
import time

# 设置中文字体（解决乱码）
matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 黑体
matplotlib.rcParams['axes.unicode_minus'] = False    # 负号显示正常

PAGE_READWRITE = 0x04
FILE_MAP_READ = 0x0004
INVALID_HANDLE_VALUE = -1

kernel32 = ctypes.WinDLL('kernel32', use_last_error=True)

OpenFileMapping = kernel32.OpenFileMappingW
OpenFileMapping.argtypes = [wintypes.DWORD, wintypes.BOOL, wintypes.LPCWSTR]
OpenFileMapping.restype = wintypes.HANDLE

MapViewOfFile = kernel32.MapViewOfFile
MapViewOfFile.argtypes = [wintypes.HANDLE, wintypes.DWORD,
                          wintypes.DWORD, wintypes.DWORD, ctypes.c_size_t]
MapViewOfFile.restype = wintypes.LPVOID

UnmapViewOfFile = kernel32.UnmapViewOfFile
UnmapViewOfFile.argtypes = [wintypes.LPCVOID]
UnmapViewOfFile.restype = wintypes.BOOL

CloseHandle = kernel32.CloseHandle
CloseHandle.argtypes = [wintypes.HANDLE]
CloseHandle.restype = wintypes.BOOL
def read_shared_memory(map_name="Local\\MySharedMemory", rows=60, cols=2, retries=10, retry_interval=1):
    size = rows * cols * 8  # 每个double占8字节
    hMap = None

    for attempt in range(retries):
        hMap = OpenFileMapping(FILE_MAP_READ, False, map_name)
        if hMap:
            break
        time.sleep(retry_interval)

    if not hMap:
        raise RuntimeError("无法打开共享内存")

    pBuf = MapViewOfFile(hMap, FILE_MAP_READ, 0, 0, size)
    if not pBuf:
        CloseHandle(hMap)
        raise ctypes.WinError(ctypes.get_last_error())

    try:
        buffer_type = ctypes.c_char * size
        buffer = buffer_type.from_address(pBuf)
        read_bytes = bytes(buffer[:])
        values = struct.unpack('d' * rows * cols, read_bytes)
        matrix = [values[i*cols:(i+1)*cols] for i in range(rows)]
        return matrix
    finally:
        UnmapViewOfFile(pBuf)
        CloseHandle(hMap)
data = read_shared_memory()
# 转换共享内存数据为x,y
x = np.array([row[0] for row in data if row[0] != 0 and row[1] != 0])
y = np.array([row[1] for row in data if row[0] != 0 and row[1] != 0])

# 线性拟合
coefficients = np.polyfit(x, y, 1)
slope, intercept = coefficients
fitted_y = slope * x + intercept

# 计算R^2
ss_res = np.sum((y - fitted_y) ** 2)
ss_tot = np.sum((y - np.mean(y)) ** 2)
r_squared = 1 - ss_res / ss_tot

# 创建图形
fig, ax = plt.subplots(figsize=(6, 4))
ax.scatter(x, y, color='blue', label='原始数据')
ax.plot(x, fitted_y, color='red', label=f'拟合线: y = {slope:.2f}x + {intercept:.2f}')
ax.set_title(f'拟合曲线：y = {slope:.2f}x + {intercept:.2f}', pad=15)

# 设置坐标轴
ax.set_xlim(0, 8)
ax.set_ylim(-60, 140)
ax.set_xticks(np.arange(1, 9, 1))  # 1 到 8
ax.set_yticks(np.arange(-60, 141, 20))

# 移除上/右边框，移动左/下边框到原点
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
ax.spines['left'].set_position('zero')
ax.spines['bottom'].set_position('zero')

# 添加箭头风格坐标轴
ax.annotate('', xy=(8, 0), xytext=(0, 0),
            arrowprops=dict(arrowstyle='->', lw=1.5, color='black'))
ax.annotate('', xy=(0, 140), xytext=(0, 0),
            arrowprops=dict(arrowstyle='->', lw=1.5, color='black'))

# 保存图像到字节流
img_bytes = io.BytesIO()
plt.savefig(img_bytes, format='png', dpi=300, bbox_inches='tight')
img_bytes.seek(0)

# # 把图像插入 Word 文档
# doc = Document()
# doc.add_heading('去除X轴0刻度的图', level=1)
# doc.add_paragraph('图中X轴从0开始，但不显示0刻度，仍保留箭头和线性拟合线：')
# doc.add_picture(img_bytes, width=Inches(5.0))
# doc.save('你干嘛.docx')

# 关闭图形
plt.close()

# ---------- 复制图像到剪贴板（仅限 Windows） ----------
def send_to_clipboard(im):
    output = io.BytesIO()
    im.convert('RGB').save(output, 'BMP')
    data = output.getvalue()[14:]  # 去掉BMP头的前14字节
    output.close()

    win32clipboard.OpenClipboard()
    win32clipboard.EmptyClipboard()
    win32clipboard.SetClipboardData(CF_DIB, data)
    win32clipboard.CloseClipboard()

# 从字节流加载为 PIL 图像并复制到剪贴板
image = Image.open(img_bytes)
send_to_clipboard(image)

# 打开 Word 文档
# os.startfile('你干嘛.docx')

import mmap
import struct
# 写入 slope 到共享内存：Local\\tempSharedMemory
slope_mem_name = "Local\\tempSharedMemory"
slope_bytes = struct.pack('3d', slope,intercept,r_squared)  # 8 字节 double

with mmap.mmap(-1, len(slope_bytes), tagname=slope_mem_name, access=mmap.ACCESS_WRITE) as slope_mmap:
    slope_mmap.write(slope_bytes)

